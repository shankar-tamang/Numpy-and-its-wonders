from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set narrow margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Header Section
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run("STUDENT’S ATTENDANCE FORM")
run.bold = True
run.font.size = Pt(14)

instructions = doc.add_paragraph(
    "The college would really be grateful if the field supervisor could fill in this attendance form for the whole "
    "Professional Placement Year duration. This form will need to be verified and enclosed together with the "
    "Confidential Field Supervisor’s Report Form. Student under this module will be allowed to take leaves only upon "
    "the discretion of the organization."
)

codes = doc.add_paragraph()
codes.add_run("Please fill in the form using the following codes:\n").bold = True
codes.add_run("P - Present\t\t\t\tPH - Public Holiday\nA - Absent\t\t\t\tTR - Outside training/ course\nMC - On Medical Leave*")

# Student Details
p = doc.add_paragraph()
p.add_run(f"Student’s Name: Shankar Tamang\nBCU ID : 24128436").bold = True

def add_month_table(doc, month_name, year, data_list):
    doc.add_paragraph(f"MONTH : {month_name} \t\t\t\t\t YEAR : {year}").bold = True
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    headers = ["DATE", "DAY", "CODE"] * 3
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fill data (3 columns of 11 rows approx)
    rows_needed = 11
    for r in range(rows_needed):
        row_cells = table.add_row().cells
        # Block 1 (Days 1-11)
        for block in range(3):
            idx = r + (block * 11)
            if idx < len(data_list):
                d, day, code = data_list[idx]
                row_cells[block*3].text = str(d)
                row_cells[block*3+1].text = day
                row_cells[block*3+2].text = code
                # Center text
                for i in range(3):
                    row_cells[block*3+i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# February 2026 Data
feb_data = [
    (1, "Sun", "P"), (2, "Mon", "P"), (3, "Tue", "P"), (4, "Wed", "P"), (5, "Thu", "P"), (6, "Fri", "PH"), (7, "Sat", "PH"),
    (8, "Sun", "P"), (9, "Mon", "P"), (10, "Tue", "P"), (11, "Wed", "P"), (12, "Thu", "P"), (13, "Fri", "PH"), (14, "Sat", "PH"),
    (15, "Sun", "PH"), # Maha Shivaratri
    (16, "Mon", "P"), (17, "Tue", "P"), (18, "Wed", "P"), (19, "Thu", "PH"), # Democracy Day
    (20, "Fri", "PH"), (21, "Sat", "PH"),
    (22, "Sun", "A"), # Personal Leave
    (23, "Mon", "P"), (24, "Tue", "P"), (25, "Wed", "P"), (26, "Thu", "P"), (27, "Fri", "PH"), (28, "Sat", "PH")
]

# March 2026 Data
mar_data = [
    (1, "Sun", "P"), (2, "Mon", "P"), (3, "Tue", "P"), (4, "Wed", "P"), (5, "Thu", "P"), (6, "Fri", "PH"), (7, "Sat", "PH"),
    (8, "Sun", "PH"), # Nari Diwas
    (9, "Mon", "P"), (10, "Tue", "P"), (11, "Wed", "P"), (12, "Thu", "P"), (13, "Fri", "PH"), (14, "Sat", "PH"),
    (15, "Sun", "P"), (16, "Mon", "P"), (17, "Tue", "P"), (18, "Wed", "P"), (19, "Thu", "P"), (20, "Fri", "PH"), (21, "Sat", "PH"),
    (22, "Sun", "P"), (23, "Mon", "P"), (24, "Tue", "P"), (25, "Wed", "P"), 
    (26, "Thu", "PH"), # Requested holiday
    (27, "Fri", "PH"), (28, "Sat", "PH"), (29, "Sun", "P"), (30, "Mon", "P"), (31, "Tue", "P")
]

add_month_table(doc, "February", "2026", feb_data)
doc.add_paragraph("\n")
add_month_table(doc, "March", "2026", mar_data)

# Footer
doc.add_paragraph("\nField supervisor’s verification:")
doc.add_paragraph("Name: Nirjal Shrestha\t\tPosition: Business Analyst")
doc.add_paragraph("Signature: ____________________\tDate: ________________")

# Save
doc.save("/mnt/data/Attendance_Form_Shankar_Tamang_Feb_March_2026.docx")